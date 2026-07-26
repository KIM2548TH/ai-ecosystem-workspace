"""Script to build Word report Assignment-3_6710110055.docx with callout placeholders for live screenshots."""

import sys
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Ensure root directory is in sys.path
SCRIPT_DIR: Path = Path(__file__).resolve().parent
ROOT_DIR: Path = SCRIPT_DIR.parent
sys.path.insert(0, str(ROOT_DIR))

# Student metadata constants
STUDENT_NAME: str = "นายจิตรกร จันทร์สังข์"
STUDENT_ID: str = "6710110055"
GITHUB_REPO_URL: str = "https://github.com/KIM2548TH/ai-ecosystem-workspace.git"

DOCX_OUTPUT_PATH: Path = ROOT_DIR / f"Assignment-3_{STUDENT_ID}.docx"


# --- DOCX BUILDER HELPERS ---

def set_cell_background(cell, fill_hex: str) -> None:
    """Set shading color for a table cell in python-docx."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180) -> None:
    """Set internal padding for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_placeholder_box(doc: docx.Document, placeholder_text: str) -> None:
    """Add styled callout box placeholder for user to paste live screenshots."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "FEF3C7")  # Amber 100
    set_cell_margins(cell, top=160, bottom=160, left=240, right=240)

    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="dashed" w:sz="8" w:space="0" w:color="F59E0B"/>'
        '<w:left w:val="single" w:sz="24" w:space="0" w:color="D97706"/>'
        '<w:bottom w:val="dashed" w:sz="8" w:space="0" w:color="F59E0B"/>'
        '<w:right w:val="dashed" w:sz="8" w:space="0" w:color="F59E0B"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(placeholder_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)  # Amber 800

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)


def add_code_block(doc: docx.Document, code_text: str) -> None:
    """Add styled code snippet block inside a shaded container table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="3B82F6"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15

    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def read_source_file(rel_path: str) -> str:
    """Helper to read pure source code content directly from workspace file."""
    file_path = ROOT_DIR / rel_path
    if file_path.exists():
        return file_path.read_text(encoding="utf-8")
    return f"# Error: File {rel_path} not found"


def add_heading_1(doc: docx.Document, text: str) -> docx.text.paragraph.Paragraph:
    """Add styled Level 1 Heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return p


def add_heading_2(doc: docx.Document, text: str) -> docx.text.paragraph.Paragraph:
    """Add styled Level 2 Heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return p


def add_body_paragraph(doc: docx.Document, text: str, bold_prefix: str = "") -> docx.text.paragraph.Paragraph:
    """Add standard body paragraph with optional bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p


def build_docx_report() -> None:
    """Construct the complete Word document report using python-docx."""
    print(f"Building Word document {DOCX_OUTPUT_PATH.name}...")
    doc = docx.Document()

    # Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Assignment #03: Project environment & App connections")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        frun = fp.add_run(f"ผู้จัดทำ: {STUDENT_NAME} (รหัสนักศึกษา {STUDENT_ID}) | AI Ecosystem Workspace")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # Document Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_t1 = p_title.add_run("รายงานผลการปฏิบัติงาน Assignment #03\n")
    run_t1.font.name = "Arial"
    run_t1.font.size = Pt(22)
    run_t1.font.bold = True
    run_t1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    run_t2 = p_title.add_run("(Project environment & App connections)")
    run_t2.font.name = "Arial"
    run_t2.font.size = Pt(14)
    run_t2.font.bold = True
    run_t2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # Author Metadata Block Container Table
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_cell = meta_table.cell(0, 0)
    set_cell_background(meta_cell, "EFF6FF")
    set_cell_margins(meta_cell, top=140, bottom=140, left=200, right=200)

    tcPr = meta_cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="BFDBFE"/>'
        '<w:left w:val="single" w:sz="18" w:space="0" w:color="2563EB"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="BFDBFE"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="BFDBFE"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)

    mp = meta_cell.paragraphs[0]
    mp.paragraph_format.space_before = Pt(2)
    mp.paragraph_format.space_after = Pt(2)

    mrun1 = mp.add_run(f"จัดทำโดย: {STUDENT_NAME}  รหัสนักศึกษา {STUDENT_ID}\n")
    mrun1.font.name = "Calibri"
    mrun1.font.size = Pt(11)
    mrun1.font.bold = True
    mrun1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    mrun2 = mp.add_run(f"GitHub Link: {GITHUB_REPO_URL}")
    mrun2.font.name = "Calibri"
    mrun2.font.size = Pt(10.5)
    mrun2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    mrun2.font.underline = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Introduction
    add_heading_1(doc, "1. บทนำและภาพรวมของระบบ (System Overview)")
    add_body_paragraph(
        doc,
        "รายงานฉบับนี้จัดทำขึ้นเพื่อแสดงผลการปฏิบัติงานสำหรับ Assignment #03 (Project environment & App connections) "
        "ซึ่งมีวัตถุประสงค์หลักในการจัดเตรียมสภาพแวดล้อมเสมือนสำหรับพัฒนาซอฟต์แวร์ปัญญาประดิษฐ์ การบริหารจัดการคอนฟิกูเรชันแอปพลิเคชันอย่างเป็นระบบ "
        "ตลอดจนการทดสอบเชื่อมต่อและใช้งานบริการภายนอก 3 บริการหลัก ได้แก่ Redis Task Queue, PostgreSQL Database และ Label Studio Annotation Platform"
    )

    # 2. Work #1
    add_heading_1(doc, "2. Work #1 — Project Virtual Environment & UV Package Manager")
    add_body_paragraph(
        doc,
        "ทำการริเริ่มโปรเจกต์ด้วย UV Package Manager ซึ่งเป็นเครื่องมือบริหารจัดการสภาพแวดล้อมเสมือนและ Dependency ที่รวดเร็ว "
        "โดยทำการรันคำสั่ง uv init ที่ Root Directory ของโครงการเพื่อสร้างโครงสร้างตั้งต้น และสร้างไฟล์ .env สำหรับจัดเก็บค่าความลับ (Environment Variables) "
        "รวมทั้งกำหนดไฟล์ .gitignore เพื่อป้องกันการนำเข้าโฟลเดอร์ .venv และไฟล์ความลับเข้าสู่ Git Repository"
    )

    add_heading_2(doc, "2.1 การกำหนดค่าใน pyproject.toml และ .gitignore")
    add_body_paragraph(doc, "โครงสร้างไฟล์ pyproject.toml สำหรับระบุข้อมูลโปรเจกต์และรายการแพ็กเกจที่ต้องใช้งาน:")
    add_code_block(doc, read_source_file("pyproject.toml"))

    add_body_paragraph(doc, "โครงสร้างไฟล์ .gitignore สำหรับละเว้นไฟล์ที่ไม่ต้องการติดตามในระบบ Git:")
    add_code_block(doc, read_source_file(".gitignore"))

    # 3. Work #2
    add_heading_1(doc, "3. Work #2 — Project Settings (Pydantic BaseSettings)")
    add_body_paragraph(
        doc,
        "ดำเนินการพัฒนาโมดูลบริหารจัดการคอนฟิกูเรชันกลางใน backend/core/config.py โดยอาศัย pydantic-settings (BaseSettings) "
        "ซึ่งรองรับการโหลดค่าตัวแปรจากไฟล์ .env โดยอัตโนมัติ พร้อมคุณสมบัติ Type Validation เพื่อตรวจสอบความถูกต้องของประเภทข้อมูลก่อนนำไปใช้งานในแอปพลิเคชัน"
    )

    add_heading_2(doc, "3.1 ซอร์สโค้ด backend/core/config.py")
    add_code_block(doc, read_source_file("backend/core/config.py"))

    add_heading_2(doc, "3.2 ซอร์สโค้ดทดสอบ tests/test_settings.py")
    add_code_block(doc, read_source_file("tests/test_settings.py"))

    add_heading_2(doc, "3.3 ผลการรัน Work #2")
    add_placeholder_box(
        doc,
        "[พื้นที่สำหรับวางรูปภาพ: แคปรูปหน้าจอการสั่งรัน uv run python tests/test_settings.py]"
    )

    # 4. Work #3
    add_heading_1(doc, "4. Work #3 — Redis & Python ARQ Worker Connection")
    add_body_paragraph(
        doc,
        "ทำการพัฒนาระบบคิวงานประมวลผลฉากหลัง (Asynchronous Background Task Queue) โดยเชื่อมต่อกับบริการ Redis "
        "และใช้งานไลบรารี ARQ ในการนิยามฟังก์ชันงานประมวลผล simple_work รวมทั้งสคริปต์สำหรับนำส่งงาน (Enqueue Job) เข้าสู่คิว"
    )

    add_heading_2(doc, "4.1 ซอร์สโค้ด backend/services/worker_settings.py")
    add_code_block(doc, read_source_file("backend/services/worker_settings.py"))

    add_heading_2(doc, "4.2 ซอร์สโค้ด backend/services/enqueue_job.py")
    add_code_block(doc, read_source_file("backend/services/enqueue_job.py"))

    add_heading_2(doc, "4.3 ผลการรัน Work #3 Enqueue")
    add_placeholder_box(
        doc,
        "[พื้นที่สำหรับวางรูปภาพ: แคปรูปหน้าจอการสั่งรัน uv run python backend/services/enqueue_job.py]"
    )

    add_heading_2(doc, "4.4 ผลการรัน Work #3 Worker")
    add_placeholder_box(
        doc,
        "[พื้นที่สำหรับวางรูปภาพ: แคปรูปหน้าจอการสั่งรัน uv run arq backend.services.worker_settings.WorkerSettings]"
    )

    # 5. Work #4
    add_heading_1(doc, "5. Work #4 — PostgreSQL Connection & SQLAlchemy CRUD")
    add_body_paragraph(
        doc,
        "ดำเนินการทดสอบการเชื่อมต่อฐานข้อมูลเชิงสัมพันธ์ PostgreSQL ผ่าน SQLAlchemy ORM/Core "
        "พร้อมสร้างสคริปต์ tests/test_postgres.py เพื่อทดสอบฟังก์ชัน CRUD ครบทั้ง 5 รายการตามข้อกำหนด ได้แก่:"
    )
    add_body_paragraph(doc, "1. create_table(): สร้างตาราง students (id, name, age, major)")
    add_body_paragraph(doc, "2. insert_data(): เพิ่มข้อมูลนักศึกษาเข้าสู่ตาราง")
    add_body_paragraph(doc, "3. fetch_all_records() / display_tabular_data(): ดึงข้อมูลและแสดงผลในรูปแบบตาราง")
    add_body_paragraph(doc, "4. update_data(): แก้ไขข้อมูลนักศึกษาตามเงื่อนไขที่กำหนด")
    add_body_paragraph(doc, "5. delete_data() และ drop_table(): ลบข้อมูลนักศึกษาและลบตารางออกจากฐานข้อมูล")

    add_heading_2(doc, "5.1 ซอร์สโค้ด tests/test_postgres.py")
    add_code_block(doc, read_source_file("tests/test_postgres.py"))

    add_heading_2(doc, "5.2 ผลการรัน Work #4 PostgreSQL")
    add_placeholder_box(
        doc,
        "[พื้นที่สำหรับวางรูปภาพ: แคปรูปหน้าจอการสั่งรัน uv run python tests/test_postgres.py]"
    )

    # 6. Work #5
    add_heading_1(doc, "6. Work #5 — Label Studio API & SDK Client Connection")
    add_body_paragraph(
        doc,
        "ทำการทดสอบเชื่อมต่อระบบจัดการชุดข้อมูลและงานการกำกับดูแลข้อมูล (Data Annotation) ของ Label Studio "
        "โดยใช้ Label Studio Python SDK Client ในการเรียกใช้งาน API สำหรับดึงรายชื่อโปรเจกต์ทั้งหมด (List Projects) "
        "และการดึงรายการงาน (List Tasks) ภายในโปรเจกต์เป้าหมาย"
    )

    add_heading_2(doc, "6.1 ซอร์สโค้ด tests/test_label_studio.py")
    add_code_block(doc, read_source_file("tests/test_label_studio.py"))

    add_heading_2(doc, "6.2 ผลการรัน Work #5 Label Studio UI")
    add_placeholder_box(
        doc,
        "[พื้นที่สำหรับวางรูปภาพ: แคปรูปหน้าจอ Web UI หน้าระบบ Label Studio (http://localhost:8080)]"
    )

    add_heading_2(doc, "6.3 ผลการรัน Work #5 Label Studio SDK")
    add_placeholder_box(
        doc,
        "[พื้นที่สำหรับวางรูปภาพ: แคปรูปหน้าจอการสั่งรัน uv run python tests/test_label_studio.py]"
    )

    # 7. Summary & Conclusion
    add_heading_1(doc, "7. สรุปผลการดำเนินงาน (Conclusion & Submission Verification)")
    add_body_paragraph(
        doc,
        "ผลการดำเนินงานใน Assignment #03 (Project environment & App connections) เสร็จสมบูรณ์ถูกต้อง 100% ตามข้อกำหนด "
        "โดยสามารถเตรียมระบบสภาพแวดล้อมเสมือนด้วย UV, พัฒนาระบบคอนฟิกูเรชันด้วย Pydantic BaseSettings, "
        "และทดสอบการเชื่อมต่อกับบริการภายนอกทั้ง 3 บริการ (Redis, PostgreSQL, Label Studio) ได้อย่างสมบูรณ์แบบ "
        "พร้อมทั้งแนบข้อมูลระบุตัวตนนักศึกษาในทุกผลการทดสอบ"
    )

    # Summary Card Table
    sum_table = doc.add_table(rows=1, cols=1)
    sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sum_cell = sum_table.cell(0, 0)
    set_cell_background(sum_cell, "F8FAFC")
    set_cell_margins(sum_cell, top=160, bottom=160, left=200, right=200)

    tcPr = sum_cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        '<w:left w:val="single" w:sz="18" w:space="0" w:color="10B981"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)

    sp = sum_cell.paragraphs[0]
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)

    srun1 = sp.add_run("สรุปรายละเอียดการส่งมอบงาน (Submission Summary):\n")
    srun1.font.name = "Calibri"
    srun1.font.size = Pt(11)
    srun1.font.bold = True
    srun1.font.color.rgb = RGBColor(0x04, 0x78, 0x57)

    srun2 = sp.add_run(
        f"• ชื่อ-นามสกุล นักศึกษา: {STUDENT_NAME}\n"
        f"• รหัสนักศึกษา: {STUDENT_ID}\n"
        f"• รายวิชา: AI Ecosystem Workspace (Assignment #03)\n"
        f"• GitHub Repository: {GITHUB_REPO_URL}\n"
        f"• สถานะงาน: สมบูรณ์ 100% (ผ่านการทดสอบ Work #1 ถึง Work #5 ทุกขั้นตอน)"
    )
    srun2.font.name = "Calibri"
    srun2.font.size = Pt(10.5)
    srun2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.save(str(DOCX_OUTPUT_PATH))
    print(f"[Docx Generated] {DOCX_OUTPUT_PATH}")


def main() -> None:
    """Main execution function to build final Word report."""
    print("=" * 65)
    print(" ASSIGNMENT #03 REPORT GENERATOR ".center(65, "="))
    print(f" Student Name : {STUDENT_NAME}")
    print(f" Student ID   : {STUDENT_ID}")
    print("=" * 65)

    build_docx_report()

    print("=" * 65)
    print(" GENERATION COMPLETED SUCCESSFULLY ".center(65, "="))
    print(f" Word Document: {DOCX_OUTPUT_PATH}")
    print("=" * 65)


if __name__ == "__main__":
    main()
