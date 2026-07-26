import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_image_placeholder(doc, title, command, image_desc):
    # Add a table with 1 row, 1 column for the border
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    
    # Set borders
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000"},
        bottom={"sz": 12, "val": "single", "color": "000000"},
        left={"sz": 12, "val": "single", "color": "000000"},
        right={"sz": 12, "val": "single", "color": "000000"},
    )
    
    # Add content to cell
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(f"[{title}]")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = RGBColor(0, 51, 102)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(f"📌 คำสั่งที่ต้องรัน/วิธีทำ: {command}")
    run2.font.size = Pt(10)
    
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run3 = p3.add_run(f"🖼️ รูปที่ต้องนำมาใส่: {image_desc}")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph() # spacing

def main():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("รายงานผลการปฏิบัติงาน Assignment #04\n(MinIO Object Storage & System Logging Architecture)")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Header Table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Student Name:"
    table.cell(0, 1).text = "นายจิตรกร จันทร์สังข์"
    table.cell(1, 0).text = "ID:"
    table.cell(1, 1).text = "6710110055"
    table.cell(2, 0).text = "GitHub:"
    table.cell(2, 1).text = "https://github.com/KIM2548TH/ai-ecosystem-workspace"
    
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    
    doc.add_paragraph()
    
    # Section 1
    doc.add_heading('ส่วนที่ 1: การติดตั้ง MinIO Service และโครงสร้างพื้นฐาน', level=1)
    doc.add_paragraph('การติดตั้ง MinIO ใน compose.yml มีการกำหนด service สำหรับ object storage โดยตั้งค่า ports สำหรับ API และ Web UI (Console) พร้อมระบุ volumes เพื่อเก็บข้อมูล')
    doc.add_paragraph('มีการแยกเก็บรหัสผ่านอย่างปลอดภัยในไฟล์ .env และให้ตัวอย่างใน .env.sample ตามหลัก Zero-Hardcoded Credentials เพื่อป้องกันไม่ให้ข้อมูลสำคัญหลุดไปใน Source Control')
    add_image_placeholder(doc, "รูปที่ 1a: หน้าจอ MinIO Console Bucket 'user-profiles'", 
                         "เข้าใช้งาน http://localhost:9001 (หรือพอร์ตที่ตั้งไว้) ล็อกอินและสร้าง/ดู bucket 'user-profiles'", 
                         "ภาพ Screenshot หน้าจอ MinIO Console ที่แสดงรายการ Bucket 'user-profiles'")
    add_image_placeholder(doc, "รูปที่ 1b: Preview รูปส่วนตัว 'จิตรกร2.jpg' บน MinIO Console", 
                         "คลิกเข้าไปที่ bucket 'user-profiles' และกดดู Preview ไฟล์ 'จิตรกร2.jpg'", 
                         "ภาพ Screenshot หน้าต่าง Preview รูปภาพจิตรกร2.jpg บน MinIO Console")
                         
    # Section 2
    doc.add_heading('ส่วนที่ 2: การอัปโหลดรูปถ่ายส่วนตัวลงใน MinIO Object Storage', level=1)
    doc.add_paragraph('การอัปโหลดสามารถทำได้ผ่าน MinIO Web Console UI และการเขียนสคริปต์ Python เพื่ออัปโหลดอัตโนมัติ')
    doc.add_paragraph('ตัวอย่างสคริปต์ (tests/minio/upload_photo.py):')
    doc.add_paragraph('''from minio import Minio
client = Minio("localhost:9000", access_key="MINIO_ROOT_USER", secret_key="MINIO_ROOT_PASSWORD", secure=False)
client.fput_object("user-profiles", "จิตรกร2.jpg", "path/to/local/จิตรกร2.jpg")''', style='Intense Quote')

    # Section 3
    doc.add_heading('ส่วนที่ 3: MinIO Python SDK (การทดสอบอัปโหลดและดาวน์โหลด)', level=1)
    doc.add_paragraph('การใช้งานสคริปต์ tests/minio/upload_download.py เพื่อทดสอบเชื่อมต่อ, ตรวจสอบ/สร้าง bucket และทดสอบ upload/download object')
    doc.add_paragraph('ตัวอย่างโค้ด (tests/minio/upload_download.py):')
    doc.add_paragraph('''if not client.bucket_exists("my-bucket"):
    client.make_bucket("my-bucket")
client.fput_object("my-bucket", "test.txt", "local_test.txt")
client.fget_object("my-bucket", "test.txt", "downloaded_test.txt")''', style='Intense Quote')
    add_image_placeholder(doc, "รูปที่ 2: ผลการรันสคริปต์ tests/minio/upload_download.py",
                         "รันคำสั่ง python tests/minio/upload_download.py ใน terminal",
                         "ภาพ Screenshot Terminal แสดงผลการทำงานสำเร็จของการอัปโหลดและดาวน์โหลดไฟล์")

    # Section 4
    doc.add_heading('ส่วนที่ 4: กลไก MinIO Object Versioning Mechanism', level=1)
    doc.add_paragraph('MinIO รองรับ Object Versioning ช่วยป้องกันการถูกเขียนทับหรือลบโดยไม่ได้ตั้งใจ สามารถเปิดใช้งานผ่าน set_bucket_versioning')
    doc.add_paragraph('ตัวอย่างโค้ด (tests/minio/versioning_test.py):')
    doc.add_paragraph('''from minio.versioningconfig import VersioningConfig, ENABLED
client.set_bucket_versioning("my-versioned-bucket", VersioningConfig(ENABLED))
# การดึงไฟล์เวอร์ชันเฉพาะ
client.fget_object("my-versioned-bucket", "data.txt", "data_v1.txt", version_id="<VERSION-ID>")''', style='Intense Quote')
    add_image_placeholder(doc, "รูปที่ 3: ผลการรันสคริปต์ tests/minio/versioning_test.py",
                         "รันคำสั่ง python tests/minio/versioning_test.py",
                         "ภาพ Screenshot Terminal แสดงประวัติ version ID ของไฟล์ที่ถูกอัปโหลดทับหลายครั้ง")

    # Section 5
    doc.add_heading('ส่วนที่ 5: สถาปัตยกรรม Custom System Logging Architecture', level=1)
    doc.add_paragraph('ระบบ Logging ใช้สถาปัตยกรรม Dual Handlers ประกอบด้วย StreamHandler (แสดงผลบน Console) และ FileHandler (บันทึกเป็นไฟล์)')
    doc.add_paragraph('โดยเก็บ Log ในรูปแบบ JSON Format แบ่งระดับ (DEBUG, INFO, WARNING, ERROR) และมีการจับสถานะ SUCCESS/FAIL พร้อมบันทึก Traceback เมื่อเกิดข้อผิดพลาด')
    doc.add_paragraph('ตัวอย่างโค้ด (utils/logger.py และ tests/logging/test_logger.py):')
    doc.add_paragraph('''@log_execution(action="test_process")
def my_process():
    log_success("Process completed", {"status": "ok"})
    log_fail("Process failed", {"error_code": 500})''', style='Intense Quote')
    add_image_placeholder(doc, "รูปที่ 4: ผลการรันสคริปต์ tests/logging/test_logger.py",
                         "รันคำสั่ง python tests/logging/test_logger.py",
                         "ภาพ Screenshot Terminal ที่เห็น JSON Log สวยงาม มีระดับของ log และสถานะต่าง ๆ")

    # Section 6
    doc.add_heading('ส่วนที่ 6: สถาปัตยกรรม Docker Container Logging System', level=1)
    doc.add_paragraph('การตั้งค่าใน compose.yml ใช้ logging driver แบบ "json-file" มีการจำกัดขนาดไฟล์เพื่อไม่ให้ disk เต็ม โดยตั้งค่า max-size: 10m และ max-file: 3')
    add_image_placeholder(doc, "รูปที่ 5: ผลการรันคำสั่ง docker compose logs",
                         "รันคำสั่ง docker compose logs -f (หรือดู log ของ service ใด service หนึ่ง)",
                         "ภาพ Screenshot Terminal แสดง logs จาก docker compose")

    # Section 7
    doc.add_heading('ส่วนที่ 7: ลิงก์ข้อมูลแหล่งประวัติโครงการ (GitHub Repository Link)', level=1)
    p = doc.add_paragraph()
    r = p.add_run("GitHub Repository: ")
    r.bold = True
    p.add_run("https://github.com/KIM2548TH/ai-ecosystem-workspace")
    
    # Save files
    paths = [
        "/home/kimbiaw/ai-eco/Assignment-4_6710110055.docx",
        "/home/kimbiaw/ai-eco/6710110055ad4.docx",
        "/home/kimbiaw/Downloads/6710110055ad4.docx"
    ]
    
    for path in paths:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        doc.save(path)
        print(f"Saved: {path}")

if __name__ == "__main__":
    main()
