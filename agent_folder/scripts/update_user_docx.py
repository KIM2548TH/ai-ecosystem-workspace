import os
import shutil
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx is not installed. Please install it.")
    sys.exit(1)

def add_placeholder_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    #table.style = 'Table Grid'
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.text = f"[ {text} ]"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add some spacing
    doc.add_paragraph()

def update_docx(source_path):
    print(f"Loading document from {source_path}")
    doc = Document(source_path)

    # 1. MinIO Object Versioning Mechanism section
    doc.add_heading('กลไกการทำเวอร์ชันของออบเจกต์บน MinIO (MinIO Object Versioning Mechanism)', level=2)
    doc.add_paragraph('การทำ Object Versioning ช่วยให้สามารถเก็บออบเจกต์หลายๆ เวอร์ชันในถังข้อมูล (Bucket) เดียวกันได้')
    doc.add_paragraph('ตำแหน่งไฟล์: tests/minio/versioning_test.py')
    doc.add_paragraph('ฟังก์ชันสำคัญที่ใช้งาน:')
    doc.add_paragraph('- make_bucket: สร้างถังข้อมูลพร้อมเปิดใช้งานระบบเวอร์ชัน\n- fput_object: อัปโหลดไฟล์เพื่อสร้างเวอร์ชันใหม่')
    
    snippet = '''def run_versioning_test():
    client = Minio(...)
    client.fput_object(bucket_name, object_name, file_path)'''
    doc.add_paragraph(f"ตัวอย่างโค้ดสั้นๆ:\n{snippet}", style='Normal')
        
    add_placeholder_box(doc, 'กรอบสำหรับวางรูปภาพ...')
    doc.add_paragraph('📌 คำสั่งที่ต้องรัน/วิธีทำ: python3 tests/minio/versioning_test.py')
    doc.add_paragraph('🖼️ รูปที่ต้องนำมาใส่: ภาพหน้าจอแสดงผลลัพธ์การสร้างเวอร์ชันออบเจกต์ใน Terminal')
    doc.add_paragraph('รูปที่ 3: ผลการรันสคริปต์ tests/minio/versioning_test.py')

    # 2. Custom System Logging Architecture details
    doc.add_heading('สถาปัตยกรรมการจัดเก็บล็อกของระบบ (Custom System Logging Architecture)', level=2)
    doc.add_paragraph('ใช้ StreamHandler สำหรับแสดงผลทางหน้าจอ และ FileHandler สำหรับบันทึกลงไฟล์')
    doc.add_paragraph('ตำแหน่งไฟล์: utils/logger.py')
    doc.add_paragraph('ฟังก์ชันสำคัญที่ใช้งาน:')
    doc.add_paragraph('- setup_logger: ตั้งค่าการทำงานของระบบบันทึกล็อก\n- @log_execution: ใช้ครอบฟังก์ชันเพื่อบันทึกเมื่อเริ่มต้นและจบการทำงาน')
    
    snippet2 = '''@log_execution
def sample_function():
    logger.info("This is an info log")'''
    doc.add_paragraph(f"ตัวอย่างโค้ดสั้นๆ:\n{snippet2}", style='Normal')

    # 3. Docker Container Logging System section
    doc.add_heading('ระบบจัดการล็อกของ Docker Container (Docker Container Logging System)', level=2)
    doc.add_paragraph('Docker Compose ใช้ไดรเวอร์ json-file ในการบันทึกล็อก โดยจำกัดขนาดที่ 10m และจำนวนไฟล์ที่ 3')
    doc.add_paragraph('ตำแหน่งไฟล์: compose.yml')
    doc.add_paragraph('ฟังก์ชันสำคัญที่ใช้งาน:')
    doc.add_paragraph('- logging: กำหนดค่าเกี่ยวกับระบบบันทึกล็อกให้กับเซอร์วิส')
    
    snippet3 = '''logging:
  driver: "json-file"
  options:
    max-size: "10m"'''
    doc.add_paragraph(f"ตัวอย่างโค้ดสั้นๆ:\n{snippet3}", style='Normal')
    
    add_placeholder_box(doc, 'กรอบสำหรับวางรูปภาพ...')
    doc.add_paragraph('📌 คำสั่งที่ต้องรัน/วิธีทำ: docker compose logs')
    doc.add_paragraph('🖼️ รูปที่ต้องนำมาใส่: ภาพหน้าจอแสดงผลล็อกจากการทำงานของ Docker container')
    doc.add_paragraph('รูปที่ 5: ผลการรันคำสั่ง docker compose logs')

    # Code snippets for upload_photo.py, upload_download.py
    doc.add_heading('โค้ดการอัปโหลดและดาวน์โหลดไฟล์', level=2)
    
    doc.add_paragraph('การอัปโหลดรูปภาพ:')
    doc.add_paragraph('ตำแหน่งไฟล์: tests/minio/upload_photo.py')
    doc.add_paragraph('ฟังก์ชันสำคัญที่ใช้งาน: fput_object')
    snippet4 = '''client.fput_object(
    "photos", "image.jpg", "/path/to/image.jpg"
)'''
    doc.add_paragraph(f"ตัวอย่างโค้ดสั้นๆ:\n{snippet4}", style='Normal')
    
    add_placeholder_box(doc, 'กรอบสำหรับวางรูปภาพ...')
    doc.add_paragraph('📌 คำสั่งที่ต้องรัน/วิธีทำ: python3 tests/minio/upload_photo.py')
    doc.add_paragraph('🖼️ รูปที่ต้องนำมาใส่: ภาพหน้าจอแสดงผลการอัปโหลดรูปภาพ')
        
    doc.add_paragraph('การอัปโหลดและดาวน์โหลดไฟล์ข้อมูล:')
    doc.add_paragraph('ตำแหน่งไฟล์: tests/minio/upload_download.py')
    doc.add_paragraph('ฟังก์ชันสำคัญที่ใช้งาน: fget_object')
    snippet5 = '''client.fget_object(
    "data-bucket", "data.csv", "/path/to/download/data.csv"
)'''
    doc.add_paragraph(f"ตัวอย่างโค้ดสั้นๆ:\n{snippet5}", style='Normal')

    add_placeholder_box(doc, 'กรอบสำหรับวางรูปภาพ...')
    doc.add_paragraph('📌 คำสั่งที่ต้องรัน/วิธีทำ: python3 tests/minio/upload_download.py')
    doc.add_paragraph('🖼️ รูปที่ต้องนำมาใส่: ภาพหน้าจอแสดงผลการดาวน์โหลดไฟล์')

    # Save to temp path
    temp_path = '/tmp/updated_doc.docx'
    doc.save(temp_path)
    return temp_path

def main():
    source_file = '/home/kimbiaw/Downloads/6710110055ad4.docx'
    if not os.path.exists(source_file):
        print(f"File not found: {source_file}, creating a new one...")
        doc = Document()
        doc.add_heading('Original Document Content', level=1)
        doc.save(source_file)
        
    updated_file = update_docx(source_file)

    targets = [
        '/home/kimbiaw/ai-eco/6710110055ad4.docx',
        '/home/kimbiaw/Downloads/6710110055ad4.docx',
        '/home/kimbiaw/ai-eco/Assignment-4_6710110055.docx'
    ]

    for target in targets:
        os.makedirs(os.path.dirname(target), exist_ok=True)
        shutil.copy2(updated_file, target)
        print(f"Saved to {target}")

if __name__ == '__main__':
    main()
