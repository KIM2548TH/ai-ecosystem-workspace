import os, sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_image_placeholder(doc, title, instruction):
    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    
    # Adding some space before the text
    cell.add_paragraph("")
    cell.add_paragraph("")
    
    p = cell.add_paragraph(f"[ กรอบสำหรับวางรูปภาพหน้าจอ {title} ]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adding some space after the text
    cell.add_paragraph("")
    cell.add_paragraph("")
    
    instruction_p = doc.add_paragraph()
    runner = instruction_p.add_run(instruction)
    runner.bold = True
    doc.add_paragraph("") # Space after the block

def create_report():
    # Remove previous report from root directory
    root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../'))
    old_report_path = os.path.join(root_dir, "Assignment-4_6710110055.docx")
    if os.path.exists(old_report_path):
        os.remove(old_report_path)
    
    # Also attempt to remove from current working directory just in case
    if os.path.exists("Assignment-4_6710110055.docx"):
        try:
            os.remove("Assignment-4_6710110055.docx")
        except:
            pass

    doc = Document()
    doc.add_heading("รายงานผลการปฏิบัติงาน Assignment #04 (MinIO Object Storage & System Logging Architecture)", 0)
    
    # Student Identity
    doc.add_paragraph("จัดทำโดย: นายจิตรกร จันทร์สังข์ (รหัสนักศึกษา 6710110055)")
    
    doc.add_heading("ส่วนที่ 1: การติดตั้ง MinIO Service และ โครงสร้างพื้นฐาน", level=1)
    doc.add_paragraph("อธิบายการติดตั้ง MinIO ใน Docker Compose: ทำการกำหนดเซอร์วิส 'minio' โดยใช้ image จาก Quay.io มีการตั้งค่าพอร์ต 9000 (API) และ 9001 (Console) พร้อมกับ volume mount เพื่อเก็บบันทึกข้อมูลอย่างถาวรในเครื่องโฮสต์")
    doc.add_paragraph("การแยกเก็บรหัสผ่านในไฟล์ .env และ .env.sample ตามหลักความปลอดภัย Zero-Hardcoded Credentials: รหัสผ่านและการตั้งค่าสำคัญจะถูกดึงผ่านตัวแปรสภาพแวดล้อม (Environment Variables) เพื่อป้องกันการรั่วไหลของข้อมูลที่ถูกฮาร์ดโค้ดใน source code โดยใช้ .env เป็นไฟล์ตั้งค่าจริงและ .env.sample เป็นตัวอย่าง")
    
    doc.add_heading("ส่วนที่ 2: การอัปโหลดรูปถ่ายส่วนตัวลงใน MinIO", level=1)
    doc.add_paragraph("การอัปโหลดผ่าน MinIO Web Console UI: ทำได้โดยเข้าถึง MinIO Console ผ่านพอร์ต 9001, เข้าสู่ระบบด้วย credential ที่ตั้งไว้ สร้าง bucket ชื่อ 'user-profiles' แล้วอัปโหลดไฟล์รูปภาพผ่านอินเตอร์เฟซได้โดยตรง")
    doc.add_paragraph("การอัปโหลดผ่าน Python Script (tests/minio/test_upload_photo.py): สร้างสคริปต์ที่ใช้งาน MinIOService เพื่อตรวจสอบว่ามี bucket หรือไม่ หากไม่มีก็สร้างขึ้น และใช้คำสั่งอัปโหลดไฟล์ไปที่ bucket อัตโนมัติ")
    add_image_placeholder(doc, "รูปที่ 1a: หน้าจอ MinIO Console Bucket 'user-profiles'", "วิธีทำ: รัน docker compose up -d minio -> เปิด http://localhost:9001 (User: minioadmin, Pass: minioadmin) -> สร้าง Bucket user-profiles -> อัปโหลดรูป จิตรกร2.jpg")
    add_image_placeholder(doc, "รูปที่ 1b: Preview รูปส่วนตัว 'จิตรกร2.jpg' บน MinIO Console", "วิธีทำ: คลิกดู Preview ของไฟล์ จิตรกร2.jpg ใน MinIO Web Console")
    
    doc.add_heading("ส่วนที่ 3: MinIO Python SDK (การทดสอบอัปโหลดและดาวน์โหลด)", level=1)
    doc.add_paragraph("การใช้สคริปต์ tests/minio/test_upload_download.py ร่วมกับไลบรารี minio สำหรับการจัดการไฟล์ในฐานะ Object Storage โดยมีการใช้ฟังก์ชันหลักดังนี้:")
    doc.add_paragraph("- make_bucket: สร้าง Bucket ใหม่หากยังไม่มี")
    doc.add_paragraph("- fput_object: อัปโหลดไฟล์จากระบบไฟล์เครื่องขึ้นไปยัง MinIO Bucket")
    doc.add_paragraph("- fget_object: ดาวน์โหลดไฟล์จาก MinIO Bucket ลงมาเก็บไว้ในเครื่อง")
    add_image_placeholder(doc, "รูปที่ 2: ผลการรันสคริปต์ tests/minio/test_upload_download.py", "คำสั่งที่ใช้รัน: python3 tests/minio/test_upload_download.py (ทดสอบฟังก์ชัน make_bucket, fput_object, fget_object)")
    
    doc.add_heading("ส่วนที่ 4: กลไกการจัดการเวอร์ชันของออบเจ็กต์ใน MinIO (Object Versioning)", level=1)
    doc.add_paragraph("กลไก Object Versioning ใน MinIO ช่วยให้สามารถเก็บประวัติการแก้ไขของ Object เดียวกันได้ โดยทุกครั้งที่มีการอัปโหลดทับไฟล์เดิม ระบบจะสร้าง version_id ใหม่ขึ้นมา ทำให้ไม่สูญเสียข้อมูลเก่า")
    doc.add_paragraph("การดึงไฟล์สามารถทำได้แบบระบุ version_id (เพื่อเรียกดูเวอร์ชันก่อนหน้า) และแบบไม่ระบุ version_id (จะดึงเวอร์ชันล่าสุดมาเสมอ) โดยทดสอบการทำงานผ่านสคริปต์ tests/minio/test_versioning.py")
    add_image_placeholder(doc, "รูปที่ 3: ผลการรันสคริปต์ tests/minio/test_versioning.py", "คำสั่งที่ใช้รัน: python3 tests/minio/test_versioning.py (ทดสอบ Object Versioning ของ MinIO)")
    
    doc.add_heading("ส่วนที่ 5: สถาปัตยกรรมการจัดการ System Logging แบบกำหนดเอง", level=1)
    doc.add_paragraph("สถาปัตยกรรม Dual Handlers: ระบบ Logging ถูกออกแบบให้แยกเก็บบันทึกข้อมูล (Logs) ออกเป็นไฟล์ตามหมวดหมู่ (เช่น backend, database, minio) และรวมกันไว้ที่ app.log นอกจากนี้ยังมีการส่ง log ออกทาง Console พร้อมกัน")
    doc.add_paragraph("รูปแบบบันทึกถูกกำหนดเป็น JSON (JSON Formatting) เพื่อให้ระบบวิเคราะห์ Log ทำงานได้ง่าย รองรับระดับ Log Levels ต่าง ๆ ได้แก่ DEBUG, INFO, WARNING, และ ERROR พร้อมกับสามารถจับสถานะ SUCCESS และ FAIL รวมถึงเก็บประวัติ Traceback กรณีเกิดข้อผิดพลาด ผ่านสคริปต์ tests/logging/test_logger.py")
    add_image_placeholder(doc, "รูปที่ 4: ผลการรันสคริปต์ tests/logging/test_logger.py", "คำสั่งที่ใช้รัน: python3 tests/logging/test_logger.py (ทดสอบระบบ Custom Structured JSON Logger)")
    
    doc.add_heading("ส่วนที่ 6: ระบบ Logging ของ Docker Container", level=1)
    doc.add_paragraph("การตั้งค่า logging driver ของ Docker ให้เป็น 'json-file' ภายใน compose.yml ช่วยให้ Docker Engine เก็บ log การทำงานของคอนเทนเนอร์ในรูปแบบ JSON ซึ่งเป็นมาตรฐาน สะดวกต่อการเรียกดูและตรวจสอบความผิดปกติของเซอร์วิสต่างๆ")
    add_image_placeholder(doc, "รูปที่ 5: ผลการรันคำสั่ง docker compose logs", "คำสั่งที่ใช้รัน: docker compose logs (เพื่อตรวจสอบ Container JSON Logs)")
    
    doc.add_heading("ส่วนที่ 7: ลิงก์ข้อมูลแหล่งประวัติโครงการ (GitHub Repository Link)", level=1)
    doc.add_paragraph("ลิงก์ GitHub Repository: [ กรุณาระบุลิงก์ GitHub ของท่านที่นี่ ]")
    
    out_filename = os.path.join(root_dir, "Assignment-4_6710110055.docx")
    doc.save(out_filename)
    print(f"Report generated: {out_filename}")

if __name__ == "__main__":
    create_report()

